"""
Module for Managing Local MS Word Documents

This module provides a collection of functions and classes designed to manage
MS Word files stored locally. It leverages the py2store framework to wrap local
binary file stores and integrates the python-docx library for handling the
content of MS Word documents. The module supports both retrieving full
`docx.Document` objects and extracting plain text from documents.

Main Classes:
    - AllLocalFilesDocxStore:
          A wrapper around a local binary store (derived from py2store's Files)
          that returns the content of files as `docx.Document` objects. This class does not
          filter files by their extension, so it may raise errors if non-MS Word files are encountered.
    - AllLocalFilesDocxTextStore:
          Extends AllLocalFilesDocxStore by applying the `get_text_from_docx` function,
          returning plain text extracted from each document instead of the full document object.
    - LocalDocxStore:
          Inherits from AllLocalFilesDocxStore and applies the `only_files_with_msword_extension`
          filter, ensuring that only files with valid MS Word extensions ('.doc' and '.docx') are processed.
          It returns `docx.Document` objects.
    - LocalDocxTextStore:
          Similar to LocalDocxStore, this class extends AllLocalFilesDocxTextStore and filters for
          valid MS Word files. It returns the extracted text from the documents.

Helper Functions:
    - _extension(k: str):
          Extracts the file extension from a filename (key) by splitting on dots.
    - has_msword_extension(k: str):
          Returns True if the key has a recognized MS Word extension ('.doc' or '.docx').
    - only_files_with_msword_extension(store):
          Filters the keys of a store to include only those with valid MS Word extensions.
    - _remove_docx_extension(k: str) and _add_docx_extension(k: str):
          Utility functions to remove or add the default '.docx' extension to keys.
    - paragraphs_text(doc):
          A generator function that yields the text of each paragraph in a document.
    - get_text_from_docx(doc, paragraph_sep='\n'):
          Concatenates the text from all paragraphs in a `docx.Document` object using
          the specified separator.
    - bytes_to_doc(doc_bytes):
          Converts a bytes stream to a `docx.Document` object using an in-memory buffer.

The relationships and dependencies between the main objects and helper functions are
illustrated in the following mermaid graph:

```mermaid
flowchart TD
    A[Files]
    B[AllLocalFilesDocxStore]
    C[AllLocalFilesDocxTextStore]
    D[LocalDocxStore]
    E[LocalDocxTextStore]

    A --> B
    B --> C
    B --> D
    C --> E

    subgraph Helper Functions
        F[bytes_to_doc]
        G[get_text_from_docx]
        H[has_msword_extension]
        I[only_files_with_msword_extension]
    end

    F --> B
    G --> C
    H --> I
    I --> D
    I --> E
```

"""

from typing import Mapping
from io import BytesIO
import docx  # (To install: pip install python-docx -- see https://automatetheboringstuff.com/chapter13/)

from dol import Files, wrap_kvs, filt_iter, Pipe


# --------------------------------------------------------------------------------------
# Baisc Helpers


def _extension(k: str):
    parts = k.split('.')
    if parts:
        return parts[-1]


_msword_extensions = {'doc', 'docx'}
_dflt_extension = '.docx'
_dflt_extension_len = len(_dflt_extension)


def has_msword_extension(k: str):
    return _extension(k) in _msword_extensions


def _remove_docx_extension(k: str):
    return k[:_dflt_extension_len]


def _add_docx_extension(k: str):
    return k + _dflt_extension


def extension_less_keys(store):
    """Will filter keys to only include .docx files, and not show that extension in keys.
    Note: Should not be used along with only_files_with_msword_extension;
    extension_less_keys already does its own key filtering.

    ```
    list(store)  # == ['this.docx', 'is.doc', 'an.pdf', 'example.docx']
    s = extension_less_keys(store)
    list(s)  # # == ['this', 'example']
    ```
    """
    return filt_iter(store, filt=_remove_docx_extension)


def paragraphs_text(doc):
    for para in doc.paragraphs:
        yield para.text


DFLT_PARAGRAPH_SEP = '\n'


def get_text_from_docx(doc, paragraph_sep=DFLT_PARAGRAPH_SEP):
    """Get text from docx.Document object.
    More precisely, 'text' will be the concatenation of the .text attributes of every paragraph.
    """
    return paragraph_sep.join(paragraphs_text(doc))


def bytes_to_doc(doc_bytes: bytes):
    return docx.Document(BytesIO(doc_bytes))


# --------------------------------------------------------------------------------------
# Mapping wrappers: Use as decorators to wrap Mappings to get the desired behavior


def only_files_with_msword_extension(store: Mapping):
    return filt_iter(store, filt=has_msword_extension)


with_bytes_to_doc_decoding = wrap_kvs(value_decoder=bytes_to_doc)
with_doc_to_text_decoding = wrap_kvs(value_decoder=get_text_from_docx)
with_bytes_to_text_decoding = wrap_kvs(
    value_decoder=Pipe(bytes_to_doc, get_text_from_docx)
)

# --------------------------------------------------------------------------------------
# Local stores (i.e. Mappings) that work with msword files
# Contributor: Note how the objects' functionality is entirely defined by the
# superclases and decorators used to wrap them.


@with_bytes_to_doc_decoding
class AllLocalFilesDocxStore(Files):
    """Local files store returning, as values, ``docx.document.Document`` objects.
    Note: Doesn't filter for valid msword extensions (.doc and .docx), so could raise errors.
    To filter for valid extensions, use LocalDocxStore instead.
    """


@wrap_kvs(obj_of_data=get_text_from_docx)
class AllLocalFilesDocxTextStore(AllLocalFilesDocxStore):
    """Local files store returning, as values, text extracted from the documents.
    Note: Doesn't filter for valid msword extensions (.doc and .docx), so could raise errors.
    To filter for valid extensions, use LocalDocxTextStore instead.
    """


@only_files_with_msword_extension
class LocalDocxStore(AllLocalFilesDocxStore):
    """Local files store returning, as values, docx objects.
    Note: Filters for valid msword extensions (.doc and .docx).
    To NOT filter for valid extensions, use AllLocalFilesDocxStore instead.

    >>> from msword import LocalDocxStore, test_data_dir
    >>> import docx
    >>> s = LocalDocxStore(test_data_dir)
    >>> assert {'with_doc_extension.doc', 'simple.docx'}.issubset(s)
    >>> v = s['with_doc_extension.doc']
    >>> assert isinstance(v, docx.document.Document)

    What does a ``docx.document.Document`` have to offer?
    If you really want to get into it, see here: https://python-docx.readthedocs.io/en/latest/

    Meanwhile, we'll give a few examples here as an amuse-bouche.

    >>> ddir = lambda x: set([xx for xx in dir(x) if not xx.startswith('_')])  # to see what an object has
    >>> assert ddir(v).issuperset({
    ...     'add_heading', 'add_page_break', 'add_paragraph', 'add_picture', 'add_section', 'add_table',
    ...     'core_properties', 'element', 'inline_shapes', 'paragraphs', 'part',
    ...     'save', 'sections', 'settings', 'styles', 'tables'
    ... })

    ``paragraphs`` is where the main content is, so let's have a look at what it has.

    >>> len(v.paragraphs)
    21
    >>> paragraph = v.paragraphs[0]
    >>> assert ddir(paragraph).issuperset({
    ...     'add_run', 'alignment', 'clear', 'insert_paragraph_before',
    ...     'paragraph_format', 'part', 'runs', 'style', 'text'
    ... })
    >>> paragraph.text
    'Section 1'
    >>> assert ddir(paragraph.style).issuperset({
    ...     'base_style', 'builtin', 'delete', 'element', 'font', 'hidden', 'locked', 'name', 'next_paragraph_style',
    ...     'paragraph_format', 'part', 'priority', 'quick_style', 'style_id', 'type', 'unhide_when_used'
    ... })
    >>> paragraph.style.style_id
    'Heading1'
    >>> paragraph.style.font.color.rgb
    RGBColor(0x2f, 0x54, 0x96)

    You get the point...

    If you're only interested in one particular aspect of the documents, you should your favorite
    py2store wrappers to get the store you really want. For example:

    >>> from py2store import wrap_kvs
    >>> ss = wrap_kvs(s, obj_of_data=lambda doc: [paragraph.style.style_id for paragraph in doc.paragraphs])
    >>> assert ss['with_doc_extension.doc'] == [
    ...     'Heading1', 'Normal', 'Normal', 'Heading2', 'Normal', 'Normal',
    ...     'Heading1', 'Normal', 'Normal', 'Normal', 'Normal', 'Normal', 'Normal', 'Normal',
    ...     'ListParagraph', 'ListParagraph', 'Normal', 'Normal', 'ListParagraph', 'ListParagraph', 'Normal'
    ... ]

    The most common use case is probably getting text, not styles, out of a document.
    It's so common, that we've done the wrapping for you:
    Just use the already wrapped LocalDocxTextStore store for that purpose.
    """


@only_files_with_msword_extension
class LocalDocxTextStore(AllLocalFilesDocxTextStore):
    """Local files store returning, as values, text extracted from the documents.
    Use this when you just want the text contents of the document.
    If you want more, you'll need to user LocalDocxStore with the appropriate content extractor
    (i.e. the obj_of_data function in a py2store.wrap_kvs wrapper).

    Note: Filters for valid msword extensions (.doc and .docx).
    To Note filter for valid extensions, use AllLocalFilesDocxTextStore instead.

    >>> from msword import LocalDocxTextStore, test_data_dir
    >>> import docx
    >>> s = LocalDocxTextStore(test_data_dir)
    >>> assert {'with_doc_extension.doc', 'simple.docx'}.issubset(s)
    >>> v = s['simple.docx']
    >>> assert isinstance(v, str)
    >>> print(v)
    Just a bit of text to show that is works. Another sentence.
    This is after a newline.
    <BLANKLINE>
    This is after two newlines.

    """
